import os
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

from utils import clean_text
from config import CHUNK_SIZE, CHUNK_OVERLAP


def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    text = "\n".join(full_text)

    return [Document(page_content=text, metadata={"source": file_path})]


def load_documents(data_path):
    documents = []

    for file in os.listdir(data_path):
        full_path = os.path.join(data_path, file)

        if file.endswith(".pdf"):
            loader = PyMuPDFLoader(full_path)
            documents.extend(loader.load())

        elif file.endswith(".txt"):
            loader = TextLoader(full_path, encoding="utf-8")
            documents.extend(loader.load())

        elif file.endswith(".docx"):
            documents.extend(load_docx(full_path))

    return documents


def preprocess_documents(documents):
    for doc in documents:
        doc.page_content = clean_text(doc.page_content)
    return documents


def chunk_documents(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    return splitter.split_documents(documents)